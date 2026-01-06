from __future__ import annotations

import argparse
import os
import tempfile
import time
from dataclasses import dataclass
from typing import Any

import requests


@dataclass
class KnowledgeBaseClient:
    api_url: str = "http://localhost:8001/api/v1"

    def __post_init__(self) -> None:
        self.session = requests.Session()
        self.token: str | None = None

    def login(self, username: str, password: str) -> bool:
        resp = self.session.post(f"{self.api_url}/auth/login", json={"username": username, "password": password})
        if resp.status_code != 200:
            return False
        data = resp.json()
        self.token = data["access_token"]
        self.session.headers.update({"Authorization": f"Bearer {self.token}"})
        return True

    def register(self, username: str, password: str) -> bool:
        resp = self.session.post(f"{self.api_url}/auth/register", json={"username": username, "password": password})
        if resp.status_code not in (200, 201):
            return False
        data = resp.json()
        self.token = data["access_token"]
        self.session.headers.update({"Authorization": f"Bearer {self.token}"})
        return True

    def health(self) -> dict[str, Any]:
        resp = self.session.get(f"{self.api_url}/health", timeout=15)
        resp.raise_for_status()
        return resp.json()

    def upload_document(self, file_path: str) -> dict[str, Any]:
        with open(file_path, "rb") as f:
            files = {"file": (os.path.basename(file_path), f)}
            resp = self.session.post(f"{self.api_url}/documents/upload", files=files, timeout=60)
        resp.raise_for_status()
        return resp.json()

    def list_documents(self, page: int = 1, page_size: int = 20, status_filter: str | None = None) -> dict[str, Any]:
        params = {"page": page, "page_size": page_size}
        if status_filter:
            params["status_filter"] = status_filter
        resp = self.session.get(f"{self.api_url}/documents", params=params, timeout=30)
        resp.raise_for_status()
        return resp.json()

    def confirm_document(self, document_id: int) -> dict[str, Any]:
        resp = self.session.post(f"{self.api_url}/documents/confirm/{document_id}", timeout=30)
        resp.raise_for_status()
        return resp.json()

    def get_document(self, document_id: int) -> dict[str, Any]:
        resp = self.session.get(f"{self.api_url}/documents/{document_id}", timeout=30)
        resp.raise_for_status()
        return resp.json()

    def get_document_response(self, document_id: int) -> requests.Response:
        return self.session.get(f"{self.api_url}/documents/{document_id}", timeout=30)

    def get_markdown_status(self, document_id: int) -> dict[str, Any]:
        resp = self.session.get(f"{self.api_url}/documents/{document_id}/markdown/status", timeout=30)
        resp.raise_for_status()
        return resp.json()

    def download_markdown(self, document_id: int) -> bytes:
        resp = self.session.get(f"{self.api_url}/documents/{document_id}/markdown/download", timeout=60)
        resp.raise_for_status()
        return resp.content

    def upload_markdown_text(self, document_id: int, markdown_text: str) -> dict[str, Any]:
        with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".md", delete=False) as f:
            f.write(markdown_text)
            path = f.name
        try:
            with open(path, "rb") as rf:
                files = {"file": (os.path.basename(path), rf, "text/markdown")}
                resp = self.session.post(f"{self.api_url}/documents/{document_id}/markdown/upload", files=files, timeout=60)
            resp.raise_for_status()
            return resp.json()
        finally:
            try:
                os.remove(path)
            except OSError:
                pass

    def delete_document(self, document_id: int) -> dict[str, Any]:
        resp = self.session.delete(f"{self.api_url}/documents/{document_id}", timeout=60)
        resp.raise_for_status()
        return resp.json()

    def list_chunks(self, document_id: int, page: int = 1, page_size: int = 50) -> dict[str, Any]:
        resp = self.session.get(
            f"{self.api_url}/documents/{document_id}/chunks",
            params={"page": page, "page_size": page_size},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()

    def create_chunk(self, document_id: int, content: str) -> dict[str, Any]:
        resp = self.session.post(
            f"{self.api_url}/documents/{document_id}/chunks",
            json={"content": content},
            timeout=60,
        )
        resp.raise_for_status()
        return resp.json()

    def update_chunk(self, document_id: int, chunk_id: int, content: str, sync_vector: bool = True) -> dict[str, Any]:
        resp = self.session.patch(
            f"{self.api_url}/documents/{document_id}/chunks/{chunk_id}",
            json={"content": content, "sync_vector": sync_vector},
            timeout=120,
        )
        resp.raise_for_status()
        return resp.json()

    def delete_chunk(self, document_id: int, chunk_id: int) -> dict[str, Any]:
        resp = self.session.delete(f"{self.api_url}/documents/{document_id}/chunks/{chunk_id}", timeout=60)
        resp.raise_for_status()
        return resp.json()

    def reembed_chunks(self, document_id: int, chunk_ids: list[int] | None = None) -> dict[str, Any]:
        payload: dict[str, Any] = {}
        if chunk_ids:
            payload["chunk_ids"] = chunk_ids
        resp = self.session.post(f"{self.api_url}/documents/{document_id}/chunks/reembed", json=payload, timeout=600)
        resp.raise_for_status()
        return resp.json()

    def batch_delete_documents(self, document_ids: list[int]) -> dict[str, Any]:
        resp = self.session.post(f"{self.api_url}/documents/batch-delete", json={"document_ids": document_ids}, timeout=60)
        resp.raise_for_status()
        return resp.json()

    def admin_list_users(self) -> dict[str, Any]:
        resp = self.session.get(f"{self.api_url}/admin/users", timeout=30)
        resp.raise_for_status()
        return resp.json()

    def get_pending_reviews(self) -> list[dict[str, Any]]:
        resp = self.session.get(f"{self.api_url}/review/pending", timeout=30)
        resp.raise_for_status()
        return resp.json()["documents"]

    def approve_document(self, document_id: int) -> dict[str, Any]:
        resp = self.session.post(f"{self.api_url}/review/approve/{document_id}", timeout=600)
        resp.raise_for_status()
        return resp.json()

    def reject_document(self, document_id: int, reason: str) -> dict[str, Any]:
        resp = self.session.post(f"{self.api_url}/review/reject/{document_id}", json={"reason": reason}, timeout=30)
        resp.raise_for_status()
        return resp.json()

    def query_knowledge_base(
        self, query: str, top_k: int = 5, model: str = "qwen2.5:32b", temperature: float = 0.7
    ) -> dict[str, Any]:
        resp = self.session.post(
            f"{self.api_url}/query",
            json={"query": query, "top_k": top_k, "model": model, "temperature": temperature},
            timeout=300,
        )
        resp.raise_for_status()
        return resp.json()

    def query_raw(self, payload: dict[str, Any]) -> dict[str, Any]:
        resp = self.session.post(f"{self.api_url}/query", json=payload, timeout=300)
        resp.raise_for_status()
        return resp.json()

    def run_acceptance(self, payload: dict[str, Any]) -> dict[str, Any]:
        resp = self.session.post(f"{self.api_url}/acceptance/run", json=payload, timeout=300)
        resp.raise_for_status()
        return resp.json()


def _pdf_escape(text: str) -> str:
    return (text or "").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_temp_pdf() -> str:
    # Minimal PDF with a single page and built-in Helvetica font.
    lines = [
        "RAG System PDF Test",
        "This PDF validates: upload -> confirm -> review -> index -> query",
        "Python async uses an event loop with async/await.",
    ]

    content_lines = ["BT", "/F1 12 Tf", "72 720 Td"]
    for i, line in enumerate(lines):
        if i > 0:
            content_lines.append("0 -18 Td")
        content_lines.append(f"({_pdf_escape(line)}) Tj")
    content_lines.append("ET")
    stream_data = ("\n".join(content_lines) + "\n").encode("ascii")

    objects: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>\n",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>\n",
        3: b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>\n",
        4: b"<< /Length %d >>\nstream\n%s\nendstream\n" % (len(stream_data), stream_data.rstrip(b"\n")),
        5: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n",
    }

    parts: list[bytes] = [b"%PDF-1.4\n"]
    offsets: dict[int, int] = {}
    for obj_id in range(1, 6):
        offsets[obj_id] = sum(len(p) for p in parts)
        parts.append(f"{obj_id} 0 obj\n".encode("ascii"))
        parts.append(objects[obj_id])
        if not parts[-1].endswith(b"\n"):
            parts.append(b"\n")
        parts.append(b"endobj\n")

    xref_offset = sum(len(p) for p in parts)
    size = 6
    xref_lines: list[bytes] = [b"xref\n", f"0 {size}\n".encode("ascii"), b"0000000000 65535 f \n"]
    for obj_id in range(1, size):
        xref_lines.append(f"{offsets[obj_id]:010d} 00000 n \n".encode("ascii"))
    trailer = [
        b"trailer\n",
        f"<< /Size {size} /Root 1 0 R >>\n".encode("ascii"),
        b"startxref\n",
        f"{xref_offset}\n".encode("ascii"),
        b"%%EOF\n",
    ]

    pdf_bytes = b"".join(parts + xref_lines + trailer)
    fd, path = tempfile.mkstemp(suffix=".pdf", prefix="kb_test_")
    os.close(fd)
    with open(path, "wb") as f:
        f.write(pdf_bytes)
    return path


def build_temp_docx() -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "Missing dependency: python-docx. Install it with `pip install python-docx` (or run inside backend container)."
        ) from exc

    doc = DocxDocument()
    doc.add_heading("RAG 系统测试文档", level=1)
    doc.add_paragraph("本文档用于验证：上传 -> 确认 -> 审核 -> 索引 -> 查询 的完整链路。")
    doc.add_paragraph("Python 异步编程基于事件循环，可以使用 async/await 语法。")
    doc.add_paragraph("FastAPI 支持异步路由函数，适合构建高并发 API 服务。")

    fd, path = tempfile.mkstemp(suffix=".docx", prefix="kb_test_")
    os.close(fd)
    doc.save(path)
    return path


def wait_for_status(client: KnowledgeBaseClient, document_id: int, status: str, timeout_s: int = 60) -> None:
    deadline = time.time() + timeout_s
    while time.time() < deadline:
        current = client.get_document(document_id)
        if current.get("status") == status:
            return
        time.sleep(1)
    raise TimeoutError(f"Timeout waiting for document #{document_id} to reach status={status}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Knowledge Base System SDK smoke test")
    parser.add_argument("--api-url", default=os.environ.get("KB_API_URL", "http://localhost:8001/api/v1"))
    parser.add_argument("--admin-username", default=os.environ.get("KB_ADMIN_USERNAME", "admin"))
    parser.add_argument("--admin-password", default=os.environ.get("KB_ADMIN_PASSWORD", "admin123"))
    parser.add_argument("--user-username", default=os.environ.get("KB_USER_USERNAME"))
    parser.add_argument("--user-password", default=os.environ.get("KB_USER_PASSWORD", "user123456"))
    parser.add_argument("--auto-register", action="store_true", help="Auto register user if login fails")
    parser.add_argument("--with-pdf", action="store_true", help="Also test PDF upload/index/query")
    parser.add_argument("--with-reject", action="store_true", help="Also test reject flow")
    parser.add_argument("--model", default=os.environ.get("KB_LLM_MODEL", "qwen3:latest"))
    args = parser.parse_args()

    api_url = args.api_url.rstrip("/")
    admin = KnowledgeBaseClient(api_url=api_url)
    user = KnowledgeBaseClient(api_url=api_url)

    health = admin.health()
    print("Health:", health)

    bad = KnowledgeBaseClient(api_url=api_url)
    if bad.login("bad_user", "bad_pass"):
        print("Unexpected: bad credentials login succeeded")
        return 1
    print("Bad-credentials login: OK (rejected as expected)")

    if not admin.login(args.admin_username, args.admin_password):
        print("Admin login failed")
        return 1
    print("Admin login OK")
    users = admin.admin_list_users()
    print(f"Admin list users: count={len(users.get('users') or [])}")

    user_username = args.user_username
    user_password = args.user_password
    if user_username:
        if user.login(user_username, user_password):
            print(f"User login OK: {user_username}")
        elif args.auto_register and user.register(user_username, user_password):
            print(f"User registered OK: {user_username}")
        else:
            print("User login failed (and register not enabled or failed)")
            return 1
    else:
        user_username = f"user_{int(time.time())}"
        if not user.register(user_username, user_password):
            print("User auto-register failed")
            return 1
        print(f"User registered OK: {user_username}")

    doc_path = build_temp_docx()
    try:
        upload = user.upload_document(doc_path)
        document_id = int(upload["document_id"])
        print("Upload:", upload)

        # Validate document listing includes the uploaded doc
        listing = user.list_documents(page=1, page_size=50)
        listed_ids = {int(d.get("id")) for d in (listing.get("documents") or []) if isinstance(d, dict) and d.get("id")}
        if document_id not in listed_ids:
            raise RuntimeError(f"Uploaded document #{document_id} not found in list_documents()")
        print(f"List documents OK (count={listing.get('total')})")

        # Validate Markdown upload/download endpoints (without waiting MinerU)
        marker = f"KB_SMOKE_MD_MARKER_{int(time.time())}"
        md_text = f"# Edited Markdown\n\n{marker}\n\n{upload.get('preview') or ''}\n"
        md_up = user.upload_markdown_text(document_id, md_text)
        print("Upload markdown:", md_up)
        md_down = user.download_markdown(document_id).decode("utf-8", errors="replace")
        if marker not in md_down:
            raise RuntimeError("Markdown download content mismatch (marker not found)")
        md_status = user.get_markdown_status(document_id)
        print("Markdown status:", md_status)

        confirm = user.confirm_document(document_id)
        print("Confirm:", confirm)

        pending = admin.get_pending_reviews()
        print(f"Pending reviews: {len(pending)}")

        approve = admin.approve_document(document_id)
        print("Approve:", approve)

        wait_for_status(user, document_id, status="indexed", timeout_s=120)

        preview = upload.get("preview") or ""
        if not preview.strip():
            raise RuntimeError("Empty preview; cannot run deterministic query test")

        query = user.query_knowledge_base(preview, top_k=5, model=args.model, temperature=0.7)
        print("Query answer:", query.get("answer"))
        print("Query sources:", query.get("sources"))

        # Optional rerank (should not break even if xinference not configured)
        rerank_query = user.query_raw(
            {
                "query": preview,
                "top_k": 5,
                "provider": "ollama",
                "model": args.model,
                "temperature": 0.7,
                "rerank": True,
                "rerank_provider": "xinference",
                "rerank_model": "bge-reranker-v2-m3",
            }
        )
        print("Query (rerank enabled) answer:", rerank_query.get("answer"))

        # Acceptance/audit workflow (admin)
        acceptance = admin.run_acceptance(
            {
                "report_document_id": document_id,
                "provider": "ollama",
                "model": args.model,
                "temperature": 0.2,
                "top_k": 5,
                "scope": "all",
            }
        )
        md = str(acceptance.get("report_markdown") or "")
        if not md.strip().startswith("# 验收审查报告"):
            raise RuntimeError("Acceptance report format mismatch (missing title)")
        print("Acceptance verdict:", acceptance.get("verdict"), "passed=", acceptance.get("passed"))

        # Chunk CRUD + vector sync
        chunk_listing = user.list_chunks(document_id, page=1, page_size=20)
        chunks = chunk_listing.get("chunks") or []
        if not chunks:
            raise RuntimeError("No chunks returned for indexed document")
        first_chunk = chunks[0]
        marker = f"KB_CHUNK_EDIT_{int(time.time())}"
        updated = user.update_chunk(
            document_id=document_id,
            chunk_id=int(first_chunk["id"]),
            content=(first_chunk.get("content") or "") + f"\n\n{marker}\n",
            sync_vector=True,
        )
        if not updated.get("vector_synced"):
            raise RuntimeError("Chunk update did not sync vector (expected vector_synced=true for indexed doc)")
        print("Chunk update OK:", {"chunk_id": first_chunk["id"], "vector_synced": updated.get("vector_synced")})

        created = user.create_chunk(document_id, content=f"{marker} NEW_CHUNK\nThis chunk is for CRUD test.")
        if not created.get("vector_synced"):
            raise RuntimeError("Chunk create did not sync vector (expected vector_synced=true for indexed doc)")
        new_chunk_id = int(created["chunk"]["id"])
        print("Chunk create OK:", {"chunk_id": new_chunk_id, "vector_synced": created.get("vector_synced")})

        reembed = user.reembed_chunks(document_id)
        print("Chunk reembed OK:", reembed)

        deleted = user.delete_chunk(document_id, new_chunk_id)
        print("Chunk delete OK:", deleted)
    finally:
        try:
            os.remove(doc_path)
        except OSError:
            pass

    # Multi-tenant isolation: another user must not access user's document
    other = KnowledgeBaseClient(api_url=api_url)
    other_name = f"user_other_{int(time.time())}"
    if not other.register(other_name, user_password):
        print("Other user register failed")
        return 1
    resp = other.get_document_response(document_id)
    if resp.status_code != 403:
        print(f"Unexpected isolation behavior: other user got status={resp.status_code}, body={resp.text}")
        return 1
    print("Multi-tenant isolation: OK (403 for other user)")

    if args.with_reject:
        reject_path = build_temp_docx()
        try:
            upload2 = user.upload_document(reject_path)
            doc2_id = int(upload2["document_id"])
            print("Upload (reject):", upload2)
            user.confirm_document(doc2_id)
            rejected = admin.reject_document(doc2_id, reason="自动化测试：内容不符合要求")
            print("Reject:", rejected)
            detail = user.get_document(doc2_id)
            print("Rejected status:", detail.get("status"))
        finally:
            try:
                os.remove(reject_path)
            except OSError:
                pass

    if args.with_pdf:
        pdf_path = build_temp_pdf()
        try:
            upload3 = user.upload_document(pdf_path)
            doc3_id = int(upload3["document_id"])
            print("Upload (pdf):", upload3)
            user.confirm_document(doc3_id)
            admin.approve_document(doc3_id)
            wait_for_status(user, doc3_id, status="indexed", timeout_s=120)

            preview3 = upload3.get("preview") or ""
            if not preview3.strip():
                raise RuntimeError("Empty PDF preview; PDF text extraction failed")
            query3 = user.query_knowledge_base(preview3, top_k=5, model=args.model, temperature=0.7)
            print("Query (pdf) answer:", query3.get("answer"))
            print("Query (pdf) sources:", query3.get("sources"))
        finally:
            try:
                os.remove(pdf_path)
            except OSError:
                pass

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
